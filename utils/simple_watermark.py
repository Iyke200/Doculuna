"""
Simple wrapper functions for the comprehensive watermark system.
Uses the existing WatermarkManager from utils.watermark
"""
import logging
from pathlib import Path
from io import BytesIO

logger = logging.getLogger(__name__)

# Import the comprehensive watermark system
try:
    from utils.watermark import WatermarkManager, WatermarkConfig
    WATERMARK_AVAILABLE = True
except ImportError:
    WATERMARK_AVAILABLE = False
    logger.warning("Watermark module not available")

# Initialize watermark manager
_watermark_manager = None

def _get_manager():
    """Get or create watermark manager instance."""
    global _watermark_manager
    if _watermark_manager is None and WATERMARK_AVAILABLE:
        _watermark_manager = WatermarkManager()
    return _watermark_manager

async def add_pdf_watermark_async(pdf_path: str, watermark_text: str = "Processed with DocuLuna - Upgrade for Watermark-Free"):
    """Add watermark to PDF file for free users (async)."""
    if not WATERMARK_AVAILABLE:
        logger.warning("Watermark not available, skipping PDF watermark")
        return
    
    try:
        manager = _get_manager()
        if not manager:
            logger.warning("Could not initialize watermark manager")
            return
        
        # Configure watermark for bottom of page
        config = WatermarkConfig(
            text=watermark_text,
            position="bottom-center",
            opacity=0.4,
            font_size=9,
            font_color=(128, 128, 128),
            rotation=0
        )
        
        # Read PDF file
        with open(pdf_path, 'rb') as f:
            pdf_data = f.read()
        
        # Add watermark
        watermarked_data = await manager.add_text_watermark(
            BytesIO(pdf_data),
            watermark_text,
            config,
            output_format='.pdf'
        )
        
        # Save watermarked PDF
        with open(pdf_path, 'wb') as f:
            f.write(watermarked_data)
        
        logger.info(f"Added watermark to PDF: {pdf_path}")
        
    except Exception as e:
        logger.error(f"Error adding PDF watermark: {e}", exc_info=True)

def add_pdf_watermark(pdf_path: str, watermark_text: str = "Processed with DocuLuna - Upgrade for Watermark-Free"):
    """Synchronous wrapper - use simple footer watermark for PDFs."""
    try:
        import fitz
        
        doc = fitz.open(pdf_path)
        
        for page in doc:
            text_rect = page.rect
            watermark_rect = fitz.Rect(
                text_rect.width * 0.1,
                text_rect.height * 0.95,
                text_rect.width * 0.9,
                text_rect.height * 0.98
            )
            
            page.insert_textbox(
                watermark_rect,
                watermark_text,
                fontsize=8,
                color=(0.5, 0.5, 0.5),
                align=fitz.TEXT_ALIGN_CENTER
            )
        
        doc.save(pdf_path, incremental=True, encryption=fitz.PDF_ENCRYPT_KEEP)
        doc.close()
        logger.info(f"Added watermark to PDF: {pdf_path}")
        
    except Exception as e:
        logger.error(f"Error adding PDF watermark: {e}", exc_info=True)

async def add_docx_watermark_async(docx_path: str, watermark_text: str = "Processed with DocuLuna - Upgrade for Watermark-Free"):
    """Add watermark to DOCX file for free users (async)."""
    # Run in executor since python-docx is sync
    import asyncio
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, add_docx_watermark, docx_path, watermark_text)

def add_docx_watermark(docx_path: str, watermark_text: str = "Processed with DocuLuna - Upgrade for Watermark-Free"):
    """Add watermark to DOCX file for free users (using footer)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.warning("python-docx not available, skipping DOCX watermark")
        return
    
    try:
        doc = Document(docx_path)
        
        # Add to footer of first section
        section = doc.sections[0]
        footer = section.footer
        
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
        else:
            paragraph = footer.add_paragraph()
        
        paragraph.text = watermark_text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if paragraph.runs:
            run = paragraph.runs[0]
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.italic = True
        
        doc.save(docx_path)
        logger.info(f"Added watermark to DOCX: {docx_path}")
        
    except Exception as e:
        logger.error(f"Error adding DOCX watermark: {e}", exc_info=True)
