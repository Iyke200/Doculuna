# utils/file_processor.py
"""
DocuLuna File Processing Utility

Secure file parsing, sanitization, and transformation for text, PDF, DOCX, and images.
Handles malware scanning, content extraction, metadata stripping, and format conversion.

Usage:
    processor = FileProcessor(max_file_size_mb=25)
    result = await processor.process_file(file_path, file_type='auto')
    sanitized_content = result['content']
    metadata = result['metadata']
"""

import logging
import os
import re
import hashlib
import magic
from typing import Dict, Any, Optional, List, Tuple, Union, IO
from pathlib import Path
from datetime import datetime
import asyncio
from dataclasses import dataclass, asdict
from enum import Enum
import tempfile

# File processing libraries
try:
    import PyPDF2
    from docx import Document
    import python_docx
    from PIL import Image, ImageOps
    import pytesseract
    PDF_AVAILABLE = True
    DOCX_AVAILABLE = True
    IMAGE_AVAILABLE = True
    OCR_AVAILABLE = True
except ImportError as e:
    logger.warning(f"Missing file processing dependency: {e}")
    PDF_AVAILABLE = False
    DOCX_AVAILABLE = False
    IMAGE_AVAILABLE = False
    OCR_AVAILABLE = False

# Security libraries
import clamav  # For malware scanning
from bleach import clean  # HTML sanitization
import exiftool  # Metadata stripping

# Local imports
from ..error_handler import ErrorHandler, ErrorContext, ErrorSeverity  # type: ignore
from ..stats import stats_tracker  # type: ignore

# Configure logging
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

@dataclass
class ProcessingResult:
    """Result of file processing operation."""
    content: str
    metadata: Dict[str, Any]
    file_type: str
    original_size: int
    processed_size: int
    warnings: List[str]
    errors: List[str]
    clean: bool  # Whether file passed security checks
    extracted_text: Optional[str] = None
    page_count: Optional[int] = None
    image_dimensions: Optional[Tuple[int, int]] = None
    ocr_confidence: Optional[float] = None

class FileType(Enum):
    """Supported file types."""
    UNKNOWN = "unknown"
    TEXT = "text"
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    IMAGE = "image"
    HTML = "html"

class SecurityThreat(Enum):
    """Identified security threats."""
    MALWARE = "malware"
    MACRO = "macro"
    PHISHING = "phishing"
    XSS = "xss"
    SQL_INJECTION = "sql_injection"
    SUSPICIOUS_EXTENSION = "suspicious_extension"
    INVALID_SIGNATURE = "invalid_signature"
    EMBEDDED_CODE = "embedded_code"

@dataclass
class ThreatDetection:
    """Threat detection results."""
    threats: List[SecurityThreat]
    severity: str  # "low", "medium", "high", "critical"
    description: str
    quarantine: bool  # Whether file should be quarantined
    sanitized: bool   # Whether threat was neutralized

class FileProcessor:
    """
    Secure file processing and sanitization engine.
    
    Features:
        - Multi-format parsing (PDF, DOCX, images, text)
        - Malware scanning with ClamAV integration
        - Content extraction and OCR
        - Metadata stripping and sanitization
        - XSS/HTML sanitization
        - File type verification and conversion
    
    Args:
        max_file_size_mb: Maximum file size in MB (default: 25)
        allowed_types: List of allowed MIME types (default: auto-detect)
        enable_ocr: Enable OCR for images (requires Tesseract)
        virus_scanner: ClamAV scanner instance (optional)
        error_handler: Error handler instance (optional)
    """
    
    # File type signatures and MIME mappings
    FILE_SIGNATURES = {
        FileType.PDF: [
            b'%PDF-',
            (0x25, 0x50, 0x44, 0x46)  # %PDF in hex
        ],
        FileType.DOCX: [
            b'PK\x03\x04',  # ZIP signature for DOCX
            b'[Content_Types].xml'
        ],
        FileType.DOC: [
            b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'  # OLE signature
        ],
        FileType.IMAGE: [
            b'\xff\xd8\xff',  # JPEG
            b'GIF8',         # GIF
            b'\x89PNG\r\n\x1a\n',  # PNG
            b'RIFF....WEBP'  # WebP
        ],
        FileType.TEXT: [
            b'{"', b'[{', b'<?xml', b'<!DOCTYPE'  # JSON, XML, HTML
        ]
    }
    
    # Dangerous file extensions
    DANGEROUS_EXTENSIONS = {
        '.exe', '.bat', '.cmd', '.com', '.scr', '.pif', '.vbs', '.js', 
        '.jar', '.war', '.ear', '.sh', '.php', '.asp', '.jsp', '.py',
        '.rb', '.pl', '.c', '.cpp', '.h', '.sql'
    }
    
    # Allowed MIME types (expandable)
    ALLOWED_MIME_TYPES = {
        'application/pdf': FileType.PDF,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': FileType.DOCX,
        'application/msword': FileType.DOC,
        'image/jpeg': FileType.IMAGE,
        'image/png': FileType.IMAGE,
        'image/gif': FileType.IMAGE,
        'image/webp': FileType.IMAGE,
        'text/plain': FileType.TEXT,
        'text/html': FileType.HTML,
        'application/json': FileType.TEXT
    }
    
    def __init__(
        self,
        max_file_size_mb: int = 25,
        allowed_types: Optional[List[str]] = None,
        enable_ocr: bool = True,
        virus_scanner: Optional[clamav.ClamdUnixSocket] = None,
        error_handler: Optional[ErrorHandler] = None,
        temp_dir: str = None
    ):
        self.max_file_size = max_file_size_mb * 1024 * 1024  # Convert to bytes
        self.allowed_types = allowed_types or list(self.ALLOWED_MIME_TYPES.keys())
        self.enable_ocr = enable_ocr and OCR_AVAILABLE
        self.virus_scanner = virus_scanner
        self.error_handler = error_handler or ErrorHandler()
        self.temp_dir = Path(temp_dir or tempfile.gettempdir())
        
        # Ensure temp directory exists and is secure
        self.temp_dir.mkdir(exist_ok=True)
        os.chmod(str(self.temp_dir), 0o700)
        
        # Initialize metadata stripper
        try:
            self.metadata_stripper = exiftool.ExifToolHelper()
        except Exception as e:
            logger.warning(f"ExifTool not available: {e}")
            self.metadata_stripper = None
        
        logger.info("FileProcessor initialized", extra={
            'max_size_mb': max_file_size_mb,
            'allowed_types_count': len(self.allowed_types),
            'ocr_enabled': self.enable_ocr,
            'virus_scanner': bool(self.virus_scanner),
            'exiftool_available': bool(self.metadata_stripper)
        })
    
    async def process_file(
        self,
        file_path: Union[str, Path, bytes, IO],
        file_type: str = 'auto',
        extract_text: bool = True,
        sanitize_html: bool = True,
        strip_metadata: bool = True,
        scan_for_malware: bool = True,
        perform_ocr: bool = False
    ) -> ProcessingResult:
        """
        Process file with full security and extraction pipeline.
        
        Args:
            file_path: File path, bytes, or file-like object
            file_type: Explicit file type or 'auto' for detection
            extract_text: Extract text content from documents
            sanitize_html: Sanitize HTML content (XSS protection)
            strip_metadata: Remove EXIF/metadata from files
            scan_for_malware: Scan with antivirus (if configured)
            perform_ocr: Perform OCR on images (if enabled)
            
        Returns:
            ProcessingResult with extracted content and security status
            
        Raises:
            FileProcessorError: If processing fails critically
            ValueError: If file type unsupported or corrupted
        """
        start_time = datetime.utcnow()
        file_processor_id = f"fp_{start_time.strftime('%Y%m%d_%H%M%S')}_{hash(str(file_path)) % 10000:04d}"
        
        context = ErrorContext(
            user_id=getattr(file_path, 'user_id', None),
            operation='file_processing',
            request_id=file_processor_id
        )
        
        try:
            logger.info("Starting file processing", extra={
                'processor_id': file_processor_id,
                'file_path': str(file_path)[:100],
                'detected_type': file_type,
                'extract_text': extract_text,
                'sanitize_html': sanitize_html,
                'strip_metadata': strip_metadata,
                'malware_scan': scan_for_malware,
                'ocr': perform_ocr
            })
            
            # Step 1: Input validation and file preparation
            file_info = await self._prepare_file(file_path)
            
            if not file_info['valid']:
                raise ValueError(f"Invalid file: {file_info['error']}")
            
            # Step 2: File type detection and validation
            detected_type = await self._detect_file_type(file_info['content'], file_type)
            if detected_type == FileType.UNKNOWN:
                raise ValueError(f"Unsupported file type: {file_info['mime_type']}")
            
            # Step 3: Security scanning
            threat_result = await self._scan_for_threats(
                file_info['content'],
                detected_type,
                scan_for_malware
            )
            
            if not threat_result['clean']:
                logger.warning("File rejected due to security threats", extra={
                    'processor_id': file_processor_id,
                    'threats': [t.value for t in threat_result['threats']],
                    'severity': threat_result['severity']
                })
                return ProcessingResult(
                    content="",
                    metadata={'security_status': 'rejected'},
                    file_type=detected_type.value,
                    original_size=len(file_info['content']),
                    processed_size=0,
                    warnings=[f"Security threats detected: {threat_result['description']}"],
                    errors=[],
                    clean=False
                )
            
            # Step 4: Metadata stripping
            if strip_metadata and self.metadata_stripper:
                file_info['content'] = await self._strip_metadata(file_info['content'], detected_type)
            
            # Step 5: Content extraction
            extraction_result = await self._extract_content(
                file_info['content'],
                detected_type,
                extract_text,
                perform_ocr
            )
            
            # Step 6: Content sanitization
            if sanitize_html and detected_type in [FileType.HTML, FileType.TEXT]:
                extraction_result['content'] = self._sanitize_content(extraction_result['content'])
            
            # Step 7: Final result assembly
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            result = ProcessingResult(
                content=extraction_result['content'],
                metadata={
                    'file_type': detected_type.value,
                    'original_filename': file_info['filename'],
                    'original_size_bytes': file_info['original_size'],
                    'mime_type': file_info['mime_type'],
                    'processing_time_seconds': round(processing_time, 3),
                    'security_status': 'clean',
                    'threats_detected': threat_result['threats'],
                    'page_count': extraction_result.get('page_count'),
                    'image_dimensions': extraction_result.get('image_dimensions'),
                    'ocr_performed': perform_ocr and detected_type == FileType.IMAGE,
                    'checksum': hashlib.sha256(file_info['content']).hexdigest()
                },
                file_type=detected_type.value,
                original_size=file_info['original_size'],
                processed_size=len(extraction_result['content'].encode('utf-8')),
                warnings=extraction_result.get('warnings', []),
                errors=extraction_result.get('errors', []),
                clean=True,
                extracted_text=extraction_result.get('extracted_text'),
                page_count=extraction_result.get('page_count'),
                image_dimensions=extraction_result.get('image_dimensions'),
                ocr_confidence=extraction_result.get('ocr_confidence')
            )
            
            # Track processing stats
            await stats_tracker.track_tool_usage(
                context.user_id or 0,
                f"file_processing_{detected_type.value}",
                processing_time,
                'success'
            )
            
            logger.info("File processing completed", extra={
                'processor_id': file_processor_id,
                'file_type': detected_type.value,
                'original_size_kb': round(file_info['original_size'] / 1024, 1),
                'processed_size_kb': round(result.processed_size / 1024, 1),
                'processing_time_s': processing_time,
                'threats': len(threat_result['threats']),
                'clean': result.clean,
                'page_count': result.page_count,
                'has_ocr': bool(result.ocr_confidence)
            })
            
            return result
            
        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            # Track failed processing
            await stats_tracker.track_tool_usage(
                context.user_id or 0,
                f"file_processing_failed",
                processing_time,
                'failed'
            )
            
            logger.error("File processing failed", exc_info=True, extra={
                'processor_id': file_processor_id,
                'error': str(e),
                'error_type': type(e).__name__,
                'processing_time_s': processing_time
            })
            
            # Use error handler if available
            if self.error_handler:
                await self.error_handler.handle_error(
                    e,
                    context=ErrorContext(
                        user_id=context.user_id,
                        operation=f"file_processing_{file_type}",
                        request_id=file_processor_id
                    ),
                    extra_data={
                        'file_path': str(file_path)[:100],
                        'detected_type': detected_type.value if 'detected_type' in locals() else 'unknown',
                        'file_size': getattr(file_info, 'original_size', 0) if 'file_info' in locals() else 0
                    },
                    severity=ErrorSeverity.ERROR,
                    category=ErrorCategory.SYSTEM
                )
            
            raise FileProcessorError(f"File processing failed: {str(e)}")
    
    async def _prepare_file(self, file_input: Union[str, Path, bytes, IO]) -> Dict[str, Any]:
        """Prepare file for processing with validation."""
        try:
            file_info = {
                'valid': False,
                'content': None,
                'filename': 'unknown',
                'original_size': 0,
                'mime_type': 'application/octet-stream',
                'extension': '',
                'error': None
            }
            
            # Handle different input types
            if isinstance(file_input, (str, Path)):
                file_path = Path(file_input)
                
                if not file_path.exists():
                    file_info['error'] = f"File not found: {file_path}"
                    return file_info
                
                if not file_path.is_file():
                    file_info['error'] = f"Not a regular file: {file_path}"
                    return file_info
                
                # Check file size
                stat = file_path.stat()
                if stat.st_size > self.max_file_size:
                    file_info['error'] = f"File too large: {stat.st_size} bytes (max: {self.max_file_size})"
                    return file_info
                
                # Check dangerous extensions
                extension = file_path.suffix.lower()
                if extension in self.DANGEROUS_EXTENSIONS:
                    file_info['error'] = f"Blocked dangerous extension: {extension}"
                    return file_info
                
                # Read file content
                async with aiofiles.open(file_path, 'rb') as f:
                    content = await f.read()
                
                file_info.update({
                    'valid': True,
                    'content': content,
                    'filename': file_path.name,
                    'original_size': stat.st_size,
                    'extension': extension
                })
                
            elif isinstance(file_input, bytes):
                if len(file_input) > self.max_file_size:
                    file_info['error'] = f"File too large: {len(file_input)} bytes"
                    return file_info
                
                file_info.update({
                    'valid': True,
                    'content': file_input,
                    'original_size': len(file_input)
                })
                
            elif hasattr(file_input, 'read'):
                # File-like object
                file_input.seek(0)
                content = file_input.read()
                
                if len(content) > self.max_file_size:
                    file_info['error'] = f"File too large: {len(content)} bytes"
                    return file_info
                
                file_info.update({
                    'valid': True,
                    'content': content,
                    'original_size': len(content)
                })
                
            else:
                file_info['error'] = f"Unsupported file input type: {type(file_input)}"
                return file_info
            
            # Detect MIME type
            try:
                mime_type = magic.from_buffer(file_info['content'], mime=True)
                file_info['mime_type'] = mime_type or 'application/octet-stream'
            except Exception:
                file_info['mime_type'] = 'application/octet-stream'
            
            logger.debug("File prepared for processing", extra={
                'filename': file_info['filename'],
                'size_bytes': file_info['original_size'],
                'mime_type': file_info['mime_type'],
                'extension': file_info['extension']
            })
            
            return file_info
            
        except Exception as e:
            logger.error("File preparation failed", exc_info=True, extra={
                'file_input_type': type(file_input),
                'error': str(e)
            })
            file_info['error'] = f"Preparation failed: {str(e)}"
            return file_info
    
    async def _detect_file_type(self, content: bytes, specified_type: str) -> FileType:
        """Detect file type using signatures and MIME detection."""
        try:
            if specified_type != 'auto':
                # Validate specified type
                for file_type, signatures in self.FILE_SIGNATURES.items():
                    if specified_type.lower() in [s.decode().lower() for s in signatures if isinstance(s, bytes)]:
                        return file_type
                return FileType.UNKNOWN
            
            # Auto-detection using signatures
            for file_type, signatures in self.FILE_SIGNATURES.items():
                for signature in signatures:
                    if isinstance(signature, bytes) and content.startswith(signature):
                        return file_type
                    elif isinstance(signature, tuple) and len(signature) == 4:
                        # Hex signature check
                        if all(content[i] == signature[i] for i in range(4)):
                            return file_type
            
            # Fallback to magic detection
            try:
                mime_type = magic.from_buffer(content, mime=True)
                for allowed_mime, file_type in self.ALLOWED_MIME_TYPES.items():
                    if mime_type == allowed_mime:
                        return file_type
            except Exception:
                pass
            
            logger.warning("File type detection failed", extra={
                'content_preview': content[:20].hex(),
                'first_bytes': content[:8],
                'detected_mime': magic.from_buffer(content, mime=True)
            })
            
            return FileType.UNKNOWN
            
        except Exception as e:
            logger.error("File type detection error", exc_info=True, extra={'error': str(e)})
            return FileType.UNKNOWN
    
    async def _scan_for_threats(self, content: bytes, file_type: FileType, 
                              enable_scan: bool) -> ThreatDetection:
        """Scan file for security threats."""
        threats = []
        severity = "low"
        
        try:
            # Basic signature validation
            if file_type == FileType.UNKNOWN:
                threats.append(SecurityThreat.SUSPICIOUS_EXTENSION)
                severity = "medium"
            
            # Extension-based blocking (already done in preparation)
            
            # ClamAV scan if enabled
            if enable_scan and self.virus_scanner:
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.tmp') as temp_file:
                        temp_file.write(content)
                        temp_path = temp_file.name
                    
                    try:
                        result = self.virus_scanner.scan_file(temp_path)
                        if result['stream'] != 'OK':
                            threats.append(SecurityThreat.MALWARE)
                            severity = "critical"
                            logger.warning("Malware detected", extra={
                                'clamav_result': result['stream'],
                                'file_type': file_type.value
                            })
                    finally:
                        os.unlink(temp_path)
                except Exception as scan_e:
                    logger.warning("ClamAV scan failed", exc_info=True, extra={
                        'error': str(scan_e),
                        'file_type': file_type.value
                    })
            
            # Macro detection in documents
            if file_type in [FileType.DOC, FileType.DOCX]:
                if await self._detect_macros(content):
                    threats.append(SecurityThreat.MACRO)
                    severity = "high"
            
            # XSS detection in text/HTML
            if file_type in [FileType.HTML, FileType.TEXT]:
                xss_threats = self._detect_xss_patterns(content.decode('utf-8', errors='ignore'))
                if xss_threats:
                    threats.extend([SecurityThreat.XSS] * len(xss_threats))
                    severity = "high" if severity == "low" else severity
            
            # SQL injection patterns
            sql_threats = self._detect_sql_injection(content.decode('utf-8', errors='ignore'))
            if sql_threats:
                threats.extend([SecurityThreat.SQL_INJECTION] * len(sql_threats))
                severity = "critical"
            
            # Embedded code detection
            if file_type in [FileType.PDF, FileType.DOCX]:
                if await self._detect_embedded_code(content, file_type):
                    threats.append(SecurityThreat.EMBEDDED_CODE)
                    severity = "high"
            
            # Determine if should quarantine
            quarantine = any(t in [SecurityThreat.MALWARE, SecurityThreat.SQL_INJECTION] 
                           for t in threats)
            
            detection = ThreatDetection(
                threats=threats,
                severity=severity,
                description=f"Detected {len(threats)} threat(s): {', '.join(t.value for t in threats)}",
                quarantine=quarantine,
                sanitized=False  # Will be sanitized later if needed
            )
            
            logger.debug("Threat detection completed", extra={
                'file_type': file_type.value,
                'threat_count': len(threats),
                'severity': severity,
                'quarantine': quarantine,
                'threat_details': [t.value for t in threats]
            })
            
            return detection
            
        except Exception as e:
            logger.error("Threat scanning failed", exc_info=True, extra={
                'file_type': file_type.value,
                'error': str(e)
            })
            return ThreatDetection(
                threats=[SecurityThreat.SYSTEM],
                severity="low",
                description=f"Scanning failed: {str(e)}",
                quarantine=False,
                sanitized=False
            )
    
    async def _detect_macros(self, content: bytes) -> bool:
        """Detect macros in document files."""
        try:
            if not DOC_AVAILABLE:
                return False
            
            # For DOCX files (ZIP-based)
            if content.startswith(b'PK\x03\x04'):
                try:
                    # Extract VBA project
                    with tempfile.NamedTemporaryFile(suffix='.docx') as temp_docx:
                        temp_docx.write(content)
                        temp_docx.flush()
                        
                        with zipfile.ZipFile(temp_docx.name, 'r') as zip_file:
                            if 'word/vbaProject.bin' in zip_file.namelist():
                                logger.warning("VBA macros detected", extra={
                                    'file_type': 'docx',
                                    'macro_file': 'vbaProject.bin'
                                })
                                return True
                except Exception as macro_e:
                    logger.debug("Macro detection failed", extra={'error': str(macro_e)})
                    return False
            
            # For legacy DOC files (OLE)
            if content.startswith(b'\xd0\xcf\x11\xe0'):
                # Basic OLE macro detection
                ole_signature = b'MM9'  # VBA module marker
                if ole_signature in content:
                    logger.warning("OLE macros detected", extra={'file_type': 'doc'})
                    return True
            
            return False
            
        except Exception as e:
            logger.warning("Macro detection error", extra={'error': str(e)})
            return False
    
    def _detect_xss_patterns(self, text: str) -> List[str]:
        """Detect XSS patterns in text content."""
        xss_patterns = [
            r'<script\b[^<]*(?:(?!<\/script>)<[^<]*)*<\/script>',
            r'on\w+\s*=',
            r'javascript\s*:',
            r'vbscript\s*:',
            r'data\s*:\s*(?:[^a-zA-Z0-9]|\w[^:]*:)',
            r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]',  # Control characters
            r'["\']?\s*(src|href|action|formaction)\s*=\s*["\']?\s*(javascript|vbscript|data):',
        ]
        
        detected = []
        for pattern in xss_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            if matches:
                detected.append(f"XSS pattern: {pattern}")
        
        return detected
    
    def _detect_sql_injection(self, text: str) -> List[str]:
        """Detect SQL injection patterns."""
        sql_patterns = [
            r'\b(union\s+(all\s+)?select|select\s+\*\s+from|insert\s+into|drop\s+table|drop\s+database|exec\s+xp_cmdshell)\b',
            r'1=1|--|;|/\*|\*/|\%00|benchmark\s*\(',
            r'\b(shutdown|grant|revoke|create|alter|truncate|delete)\b',
            r'@@version|sp_password|xp_cmdshell',
            r'\b(and|or)\s+\d+\s*=',
            r'0x[0-9a-f]+'  # Hex encoded payloads
        ]
        
        detected = []
        for pattern in sql_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                detected.append(f"SQL pattern: {pattern}")
        
        return detected
    
    async def _detect_embedded_code(self, content: bytes, file_type: FileType) -> bool:
        """Detect embedded executable code."""
        try:
            if file_type == FileType.PDF:
                # Check for JavaScript in PDF
                js_indicators = [b'/JS', b'/JavaScript', b'/AA', b'/OpenAction']
                return any(indicator in content for indicator in js_indicators)
            
            elif file_type == FileType.DOCX:
                # Check for embedded objects or scripts
                obj_indicators = [b'<w:embedded', b'<v:shape', b'http://schemas.openxmlformats.org']
                return any(indicator in content for indicator in obj_indicators)
            
            return False
            
        except Exception as e:
            logger.warning("Embedded code detection failed", extra={
                'file_type': file_type.value,
                'error': str(e)
            })
            return False
    
    async def _strip_metadata(self, content: bytes, file_type: FileType) -> bytes:
        """Strip metadata from files."""
        try:
            if not self.metadata_stripper:
                logger.debug("Metadata stripping skipped - ExifTool not available")
                return content
            
            async with self.metadata_stripper as et:
                # Create temporary file for processing
                with tempfile.NamedTemporaryFile(delete=False) as temp_in, \
                     tempfile.NamedTemporaryFile(delete=False) as temp_out:
                    
                    # Write input file
                    temp_in.write(content)
                    temp_in.flush()
                    os.fsync(temp_in.fileno())
                    
                    # Strip metadata
                    await et.update_tags(
                        temp_in.name,
                        {str(temp_out.name): {}},  # Output to new file
                        params=['-all=', '-o', temp_out.name]  # Remove all metadata
                    )
                    
                    # Read cleaned file
                    with open(temp_out.name, 'rb') as cleaned_file:
                        cleaned_content = cleaned_file.read()
                    
                    # Cleanup
                    os.unlink(temp_in.name)
                    os.unlink(temp_out.name)
                
                logger.debug("Metadata stripped", extra={
                    'file_type': file_type.value,
                    'original_size': len(content),
                    'cleaned_size': len(cleaned_content)
                })
                
                return cleaned_content
                
        except Exception as e:
            logger.warning("Metadata stripping failed", extra={
                'file_type': file_type.value,
                'error': str(e)
            })
            return content
    
    async def _extract_content(self, content: bytes, file_type: FileType, 
                             extract_text: bool, perform_ocr: bool) -> Dict[str, Any]:
        """Extract content based on file type."""
        extraction_result = {
            'content': '',
            'extracted_text': '',
            'page_count': None,
            'image_dimensions': None,
            'ocr_confidence': None,
            'warnings': [],
            'errors': []
        }
        
        try:
            if file_type == FileType.PDF and PDF_AVAILABLE:
                result = await self._extract_pdf_content(content)
                extraction_result.update(result)
                
            elif file_type == FileType.DOCX and DOCX_AVAILABLE:
                result = await self._extract_docx_content(content)
                extraction_result.update(result)
                
            elif file_type == FileType.DOC:
                result = await self._extract_doc_content(content)
                extraction_result.update(result)
                
            elif file_type == FileType.IMAGE and IMAGE_AVAILABLE:
                result = await self._extract_image_content(content, perform_ocr)
                extraction_result.update(result)
                
            elif file_type in [FileType.TEXT, FileType.HTML]:
                # Simple text extraction
                try:
                    text_content = content.decode('utf-8', errors='replace')
                    extraction_result['content'] = text_content
                    extraction_result['extracted_text'] = text_content
                except UnicodeDecodeError:
                    # Try different encodings
                    encodings = ['utf-8', 'latin-1', 'ascii']
                    for encoding in encodings:
                        try:
                            text_content = content.decode(encoding, errors='replace')
                            extraction_result['content'] = text_content
                            extraction_result['extracted_text'] = text_content
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        extraction_result['warnings'].append('Could not decode text content')
                        extraction_result['content'] = repr(content[:500])
                
            else:
                extraction_result['warnings'].append(f"Content extraction not supported for {file_type.value}")
                extraction_result['content'] = repr(content[:1000])
            
            logger.debug("Content extraction completed", extra={
                'file_type': file_type.value,
                'content_length': len(extraction_result['content']),
                'extracted_text_length': len(extraction_result['extracted_text']),
                'page_count': extraction_result['page_count'],
                'has_ocr': bool(extraction_result['ocr_confidence'])
            })
            
            return extraction_result
            
        except Exception as e:
            logger.error("Content extraction failed", exc_info=True, extra={
                'file_type': file_type.value,
                'error': str(e)
            })
            extraction_result['errors'].append(f"Extraction failed: {str(e)}")
            extraction_result['content'] = repr(content[:500])
            return extraction_result
    
    async def _extract_pdf_content(self, content: bytes) -> Dict[str, Any]:
        """Extract text and metadata from PDF."""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(io.BytesIO(content))
            page_count = len(reader.pages)
            
            full_text = []
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        full_text.append(f"--- Page {page_num} ---\n{text.strip()}")
                except Exception as page_e:
                    logger.warning(f"Failed to extract page {page_num}", extra={'error': str(page_e)})
                    full_text.append(f"[Page {page_num}: Extraction failed]")
            
            extracted_text = '\n\n'.join(full_text)
            
            return {
                'content': f"PDF Document ({page_count} pages)\n\n" + extracted_text,
                'extracted_text': extracted_text,
                'page_count': page_count,
                'warnings': [],
                'errors': []
            }
            
        except Exception as e:
            logger.error("PDF extraction failed", exc_info=True, extra={'error': str(e)})
            return {
                'content': "PDF processing failed",
                'extracted_text': '',
                'page_count': None,
                'warnings': [],
                'errors': [f"PDF extraction error: {str(e)}"]
            }
    
    async def _extract_docx_content(self, content: bytes) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(content))
            full_text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text:
                        full_text.append(f"Table: {row_text}")
            
            extracted_text = '\n'.join(full_text)
            
            return {
                'content': f"DOCX Document\n\n{extracted_text}",
                'extracted_text': extracted_text,
                'page_count': None,
                'warnings': [],
                'errors': []
            }
            
        except Exception as e:
            logger.error("DOCX extraction failed", exc_info=True, extra={'error': str(e)})
            return {
                'content': "DOCX processing failed",
                'extracted_text': '',
                'page_count': None,
                'warnings': [],
                'errors': [f"DOCX extraction error: {str(e)}"]
            }
    
    async def _extract_doc_content(self, content: bytes) -> Dict[str, Any]:
        """Extract text from legacy DOC file (basic implementation)."""
        try:
            # Legacy DOC files require OLE parsing - simplified text extraction
            # In production, use antiword or libreoffice --headless conversion
            
            # Basic fallback: treat as binary and extract printable text
            text_content = ''
            try:
                # Try to decode as text (may not work well for binary DOC)
                text_content = content.decode('utf-8', errors='ignore')
                text_content = re.sub(r'[^\x20-\x7E\n\r\t]', '', text_content)  # ASCII printable
            except UnicodeDecodeError:
                # Binary extraction fallback
                text_content = ''.join(chr(b) for b in content if 32 <= b <= 126)
            
            cleaned_text = re.sub(r'\s+', ' ', text_content.strip())[:10000]  # Limit size
            
            return {
                'content': f"DOC Document (Legacy)\n\n{cleaned_text}",
                'extracted_text': cleaned_text,
                'page_count': None,
                'warnings': ['Limited extraction for legacy DOC format'],
                'errors': []
            }
            
        except Exception as e:
            logger.error("DOC extraction failed", exc_info=True, extra={'error': str(e)})
            return {
                'content': "DOC processing failed",
                'extracted_text': '',
                'page_count': None,
                'warnings': [],
                'errors': [f"DOC extraction error: {str(e)}"]
            }
    
    async def _extract_image_content(self, content: bytes, 
                                   perform_ocr: bool) -> Dict[str, Any]:
        """Extract content from image files."""
        try:
            from PIL import Image
            
            image = Image.open(io.BytesIO(content))
            
            # Get image info
            width, height = image.size
            format_name = image.format or 'Unknown'
            
            content_info = f"Image: {width}x{height} pixels, {format_name} format"
            
            extracted_text = ''
            ocr_confidence = None
            
            if perform_ocr and self.enable_ocr:
                try:
                    # Convert to grayscale if needed
                    if image.mode != 'L':
                        image = image.convert('L')
                    
                    # Perform OCR
                    extracted_text = pytesseract.image_to_string(image)
                    ocr_confidence = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                    # Calculate average confidence
                    confidences = [int(conf) for conf in ocr_confidence['conf'] if int(conf) > 0]
                    ocr_confidence = sum(confidences) / len(confidences) if confidences else 0
                    
                    if not extracted_text.strip():
                        extracted_text = ''
                        ocr_confidence = None
                        
                except Exception as ocr_e:
                    logger.warning("OCR processing failed", exc_info=True, extra={'error': str(ocr_e)})
                    extraction_result['warnings'].append('OCR unavailable or failed')
            
            return {
                'content': content_info + (f"\n\nOCR Text: {extracted_text}" if extracted_text else ""),
                'extracted_text': extracted_text,
                'image_dimensions': (width, height),
                'ocr_confidence': ocr_confidence,
                'warnings': [],
                'errors': []
            }
            
        except Exception as e:
            logger.error("Image processing failed", exc_info=True, extra={'error': str(e)})
            return {
                'content': "Image processing failed",
                'extracted_text': '',
                'image_dimensions': None,
                'ocr_confidence': None,
                'warnings': [],
                'errors': [f"Image processing error: {str(e)}"]
            }
    
    def _sanitize_content(self, content: str) -> str:
        """Sanitize text content for XSS and injection attacks."""
        try:
            # Bleach HTML sanitization
            if '<' in content and '>' in content:
                # Likely HTML content
                from bleach import clean
                from bleach.sanitizer import Cleaner
                cleaner = Cleaner(
                    tags=['p', 'br', 'strong', 'em', 'ul', 'ol', 'li', 'a'],
                    attributes={'a': ['href']},
                    styles=[],
                    protocols=['http', 'https'],
                    convert_entities=True
                )
                sanitized = cleaner.clean(content)
            else:
                sanitized = content
            
            # Remove control characters
            sanitized = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', sanitized)
            
            # Neutralize common injection patterns
            injection_patterns = [
                (r'(\b(union|select|insert|delete|drop|update|create|alter|exec|sp_|xp_)\b)', '[SQL_REDACTED]'),
                (r'(on\w+\s*=)', '[EVENT_REDACTED]'),
                (r'(javascript\s*:)', '[JS_REDACTED]'),
                (r'(<script\b)', '[SCRIPT_REDACTED]')
            ]
            
            for pattern, replacement in injection_patterns:
                sanitized = re.sub(pattern, replacement, sanitized, flags=re.IGNORECASE)
            
            # Limit length for processing
            if len(sanitized) > 1000000:  # 1MB limit
                sanitized = sanitized[:1000000] + "\n[Content truncated for safety]"
                logger.warning("Content truncated during sanitization", extra={
                    'original_length': len(content),
                    'sanitized_length': len(sanitized)
                })
            
            return sanitized
            
        except Exception as e:
            logger.error("Content sanitization failed", exc_info=True, extra={'error': str(e)})
            return content[:5000]  # Truncated fallback
    
    async def _sanitize_file(self, content: bytes, threats: List[SecurityThreat], 
                           file_type: FileType) -> Tuple[bytes, bool]:
        """Apply sanitization based on detected threats."""
        sanitized = content
        fully_sanitized = True
        
        try:
            for threat in threats:
                if threat == SecurityThreat.XSS:
                    # For text/HTML files, sanitize content
                    if file_type in [FileType.HTML, FileType.TEXT]:
                        try:
                            text_content = content.decode('utf-8', errors='replace')
                            sanitized_text = self._sanitize_content(text_content)
                            sanitized = sanitized_text.encode('utf-8')
                        except Exception:
                            fully_sanitized = False
                
                elif threat == SecurityThreat.SQL_INJECTION:
                    # Remove executable content
                    if file_type in [FileType.TEXT, FileType.HTML]:
                        text_content = content.decode('utf-8', errors='replace')
                        # Remove SQL keywords and patterns
                        sql_cleaned = re.sub(
                            r'\b(union\s+(all\s+)?select|insert\s+into|drop\s+(table|database)|exec\s+\w+)\b',
                            '[REMOVED]',
                            text_content,
                            flags=re.IGNORECASE
                        )
                        sanitized = sql_cleaned.encode('utf-8')
                
                elif threat == SecurityThreat.MACRO:
                    # For document files, attempt macro removal
                    if file_type in [FileType.DOCX]:
                        sanitized = await self._remove_macros_from_docx(content)
                    elif file_type in [FileType.DOC]:
                        sanitized = await self._remove_macros_from_doc(content)
                    else:
                        fully_sanitized = False
                
                elif threat == SecurityThreat.EMBEDDED_CODE:
                    # Strip embedded scripts from PDF/DOCX
                    if file_type == FileType.PDF:
                        sanitized = await self._remove_pdf_javascript(content)
                    else:
                        fully_sanitized = False
                
                else:
                    # Unknown threat type - quarantine
                    fully_sanitized = False
            
            logger.debug("File sanitization applied", extra={
                'file_type': file_type.value,
                'threats_handled': len(threats),
                'fully_sanitized': fully_sanitized,
                'size_change': len(sanitized) - len(content)
            })
            
            return sanitized, fully_sanitized
            
        except Exception as e:
            logger.error("File sanitization failed", exc_info=True, extra={
                'file_type': file_type.value,
                'threats': [t.value for t in threats],
                'error': str(e)
            })
            return content, False
    
    async def _remove_macros_from_docx(self, content: bytes) -> bytes:
        """Remove macros from DOCX file."""
        try:
            from zipfile import ZipFile
            import xml.etree.ElementTree as ET
            
            with tempfile.NamedTemporaryFile(suffix='.docx') as temp_in, \
                 tempfile.NamedTemporaryFile(suffix='.docx') as temp_out:
                
                # Write original file
                temp_in.write(content)
                temp_in.flush()
                
                # Create clean DOCX
                with ZipFile(temp_in.name, 'r') as zin:
                    with ZipFile(temp_out.name, 'w') as zout:
                        for item in zin.infolist():
                            # Skip VBA project and macro files
                            if not any(skip in item.filename.lower() for skip in 
                                     ['vbaproject.bin', 'macros', 'vba']):
                                zout.writestr(item, zin.read(item.filename))
                
                with open(temp_out.name, 'rb') as cleaned_file:
                    return cleaned_file.read()
                    
        except Exception as e:
            logger.warning("DOCX macro removal failed", extra={'error': str(e)})
            return content
    
    async def _remove_pdf_javascript(self, content: bytes) -> bytes:
        """Remove JavaScript from PDF."""
        try:
            from PyPDF2 import PdfReader, PdfWriter
            
            reader = PdfReader(io.BytesIO(content))
            writer = PdfWriter()
            
            for page in reader.pages:
                # Remove JavaScript actions
                if '/AA' in page:
                    del page['/AA']
                if '/JS' in page:
                    del page['/JS']
                if '/JavaScript' in page:
                    del page['/JavaScript']
                
                writer.add_page(page)
            
            output = io.BytesIO()
            writer.write(output)
            return output.getvalue()
            
        except Exception as e:
            logger.warning("PDF JavaScript removal failed", extra={'error': str(e)})
            return content
    
    def _sanitize_html_content(self, html_content: str) -> str:
        """Sanitize HTML content using Bleach."""
        try:
            from bleach import clean
            from bleach.sanitizer import Cleaner
            
            # Safe tags and attributes
            cleaner = Cleaner(
                tags=['p', 'br', 'strong', 'em', 'ul', 'ol', 'li', 'a', 'h1', 'h2', 'h3', 'div'],
                attributes={
                    'a': ['href'],
                    '*: None  # No other attributes
                },
                styles=[],
                protocols=['http', 'https', 'mailto'],
                convert_entities=True
            )
            
            sanitized = cleaner.clean(html_content)
            
            # Additional regex-based cleaning
            sanitized = re.sub(r'<script\b[^<]*(?:(?!<\/script>)<[^<]*)*<\/script>', '', sanitized, flags=re.IGNORECASE)
            sanitized = re.sub(r'on\w+\s*=', '', sanitized, flags=re.IGNORECASE)
            sanitized = re.sub(r'javascript\s*:', '', sanitized, flags=re.IGNORECASE)
            
            return sanitized
            
        except Exception as e:
            logger.warning("HTML sanitization failed", extra={'error': str(e)})
            # Fallback: strip all tags
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            return soup.get_text()
    
    async def convert_file_format(self, content: bytes, source_type: FileType, 
                                target_format: str) -> Optional[bytes]:
        """Convert file to different format."""
        try:
            if source_type == FileType.PDF and target_format == 'text':
                # Extract text from PDF
                return await self._extract_pdf_text(content)
            
            elif source_type == FileType.IMAGE and target_format == 'pdf':
                # Convert image to PDF
                return await self._image_to_pdf(content)
            
            elif source_type == FileType.DOCX and target_format == 'pdf':
                # DOCX to PDF conversion (requires external tools)
                return await self._docx_to_pdf(content)
            
            else:
                logger.warning("Unsupported format conversion", extra={
                    'source': source_type.value,
                    'target': target_format
                })
                return None
                
        except Exception as e:
            logger.error("Format conversion failed", exc_info=True, extra={
                'source': source_type.value,
                'target': target_format,
                'error': str(e)
            })
            return None
    
    async def _extract_pdf_text(self, content: bytes) -> bytes:
        """Extract plain text from PDF."""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(io.BytesIO(content))
            full_text = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text.strip())
            
            text_content = '\n\n'.join(full_text)
            return text_content.encode('utf-8')
            
        except Exception as e:
            logger.error("PDF text extraction failed", extra={'error': str(e)})
            return b"PDF text extraction failed"
    
    async def _image_to_pdf(self, content: bytes) -> bytes:
        """Convert image to PDF."""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from io import BytesIO
            
            image = Image.open(BytesIO(content))
            
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            
            # Scale image to fit page
            img_width, img_height = image.size
            scale = min(width/img_width, height/img_height)
            new_width = img_width * scale
            new_height = img_height * scale
            
            # Center image
            x = (width - new_width) / 2
            y = (height - new_height) / 2
            
            c.drawImage(image.filename if hasattr(image, 'filename') else 'image', 
                       x, y, new_width, new_height)
            c.save()
            
            buffer.seek(0)
            return buffer.read()
            
        except Exception as e:
            logger.error("Image to PDF conversion failed", extra={'error': str(e)})
            return b"Image conversion failed"
    
    async def _docx_to_pdf(self, content: bytes) -> bytes:
        """Convert DOCX to PDF using LibreOffice (headless)."""
        try:
            # This requires LibreOffice command-line tools
            import subprocess
            
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_file:
                docx_file.write(content)
                docx_path = docx_file.name
            
            pdf_path = docx_path + '.pdf'
            
            # LibreOffice headless conversion
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(pdf_path),
                docx_path
            ]
            
            process = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            
            stdout, stderr = await process.communicate()
            
            if process.returncode == 0 and os.path.exists(pdf_path):
                with open(pdf_path, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                
                # Cleanup
                os.unlink(docx_path)
                os.unlink(pdf_path)
                
                return pdf_content
            else:
                logger.error("DOCX to PDF conversion failed", extra={
                    'return_code': process.returncode,
                    'stderr': stderr.decode()
                })
                return b"DOCX conversion failed"
                
        except FileNotFoundError:
            logger.warning("LibreOffice not available for DOCX conversion")
            return b"Conversion requires LibreOffice"
        except Exception as e:
            logger.error("DOCX to PDF conversion error", extra={'error': str(e)})
            return b"Conversion failed"

class FileProcessorError(Exception):
    """Custom exception for file processing errors."""
    pass

# Global processor instance (configured in main app)
file_processor: Optional[FileProcessor] = None

def initialize_file_processor(
    max_file_size_mb: int = 25,
    virus_scanner_config: Optional[Dict[str, str]] = None,
    error_handler_instance: Optional[ErrorHandler] = None
) -> FileProcessor:
    """Initialize global file processor."""
    global file_processor
    
    if file_processor is None:
        # Initialize ClamAV scanner if configured
        virus_scanner = None
        if virus_scanner_config:
            try:
                if virus_scanner_config.get('type') == 'clamav':
                    virus_scanner = clamav.ClamdUnixSocket(virus_scanner_config['socket_path'])
                logger.info("ClamAV scanner initialized")
            except Exception as e:
                logger.warning(f"Failed to initialize virus scanner: {e}")
                virus_scanner = None
        
        file_processor = FileProcessor(
            max_file_size_mb=max_file_size_mb,
            virus_scanner=virus_scanner,
            error_handler=error_handler_instance
        )
    
    return file_processor

# Convenience functions for quick file processing
async def process_document(file_input: Union[str, Path, bytes], 
                         user_id: Optional[int] = None) -> ProcessingResult:
    """Quick document processing with defaults."""
    processor = file_processor
    if not processor:
        raise RuntimeError("File processor not initialized")
    
    context = ErrorContext(user_id=user_id, operation='quick_document_processing')
    
    try:
        return await processor.process_file(
            file_input,
            file_type='auto',
            extract_text=True,
            sanitize_html=True,
            strip_metadata=True,
            scan_for_malware=True,
            perform_ocr=False
        )
    except Exception as e:
        # Use global error handler
        from .error_handler import error_handler
        await error_handler.handle_error(e, context=context)
        raise FileProcessorError(f"Document processing failed: {str(e)}")

async def extract_text_from_file(file_input: Union[str, Path, bytes]) -> str:
    """Extract text content from any supported file."""
    processor = file_processor
    if not processor:
        raise RuntimeError("File processor not initialized")
    
    result = await processor.process_file(
        file_input,
        extract_text=True,
        sanitize_html=False,
        strip_metadata=False,
        scan_for_malware=False
    )
    
    return result.extracted_text or result.content

def sanitize_file_content(content: str) -> str:
    """Quick content sanitization for user input."""
    processor = file_processor
    if not processor:
        from bleach import clean
        return clean(content, tags=[], strip=True)
    
    return processor._sanitize_content(content)
