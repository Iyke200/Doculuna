# file_handler.py - File handling with gamification integration
import logging
import os
import tempfile
import asyncio
import time
from typing import Optional
from aiogram import Dispatcher, types, Bot
from aiogram.fsm.context import FSMContext
from aiogram.utils.keyboard import InlineKeyboardBuilder
from aiogram.types import FSInputFile

from utils.usage_tracker import check_usage_limit, increment_usage
from utils.watermark import WatermarkManager, WatermarkConfig
from tools.pdf_to_word import PDFToWordConverter
from tools.word_to_pdf import WordToPDFConverter
from tools.compress import PDFCompressor, DOCXCompressor, CompressionLevel
from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from io import BytesIO

# Import gamification and history modules
from handlers.gamification import gamification_engine
from handlers.history import log_operation, get_history_count
from handlers.file_naming import sanitize_filename, generate_output_filename

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize watermark manager
watermark_manager = WatermarkManager()
WATERMARK_TEXT = "Processed with DocuLuna - Upgrade for Watermark-Free"

# XP rewards for different operations
XP_REWARDS = {
    "convert": 50,
    "compress": 40,
    "image_to_pdf": 45,
    "merge": 60,
    "split": 55,
    "ocr": 65,
}


async def process_gamification(user_id: int, operation: str, filename: str, duration: float, file_size: int = 0, output_filename: Optional[str] = None):
    """Process gamification rewards and log the operation."""
    try:
        # Update streak first
        streak_result = await gamification_engine.update_streak(user_id)
        
        # Add XP for the operation
        xp_amount = XP_REWARDS.get(operation, 50)
        xp_result = await gamification_engine.add_xp(user_id, xp_amount)
        
        # Log the operation to history
        await log_operation(
            user_id=user_id,
            operation=operation,
            filename=filename,
            duration=duration,
            status='success',
            file_size=file_size,
            output_filename=output_filename
        )
        
        # Check for history-based achievements
        history_count = await get_history_count(user_id)
        await gamification_engine.check_history_achievements(user_id, history_count)
        
        # Build reward message
        rewards = []
        if xp_result.get("leveled_up"):
            rewards.extend(xp_result.get("messages", []))
        if streak_result.get("message"):
            rewards.append(streak_result["message"])
        
        return {
            "xp_gained": xp_amount,
            "leveled_up": xp_result.get("leveled_up", False),
            "new_level": xp_result.get("new_level"),
            "new_rank": xp_result.get("new_rank"),
            "streak": streak_result.get("streak", 0),
            "streak_increased": streak_result.get("increased", False),
            "messages": rewards
        }
    except Exception as e:
        logger.error(f"Error processing gamification for user {user_id}: {e}", exc_info=True)
        return {"xp_gained": 0, "messages": []}


async def handle_document(message: types.Message, state: FSMContext):
    """Handle file received."""
    user_id = message.from_user.id
    
    can_process = await check_usage_limit(user_id)
    if not can_process:
        from config import FREE_USAGE_LIMIT
        limit_message = (
            "⚠️ Daily limit reached!\n\n"
            f"You've used all {FREE_USAGE_LIMIT} free document processing actions for today.\n\n"
            "💎 Upgrade to Premium for unlimited processing!\n"
            "Click /premium to see plans."
        )
        await message.reply(limit_message)
        return
    
    try:
        file_received_text = (
            "📄 File received successfully!\n"
            "What would you like to do with it? 👇"
        )
        
        builder = InlineKeyboardBuilder()
        builder.button(text="🔁 Convert File", callback_data="convert_file")
        builder.button(text="📊 Merge PDFs", callback_data="merge_pdfs")
        builder.button(text="✂️ Split PDF", callback_data="split_pdf")
        builder.button(text="🗜️ Compress", callback_data="compress_file")
        builder.adjust(2, 2)
        
        await message.reply(file_received_text, reply_markup=builder.as_markup())
        
        await state.update_data(file_id=message.document.file_id if message.document else None)
        await state.update_data(file_name=message.document.file_name if message.document else "file")
        
    except Exception as e:
        logger.error(f"Error handling document: {e}", exc_info=True)
        await message.reply(
            "⚠️ Oops! Something went wrong while processing your file.\n"
            "Please try again later or contact @DocuLunaSupport"
        )


async def handle_photo(message: types.Message, state: FSMContext):
    """Handle image received."""
    user_id = message.from_user.id
    
    can_process = await check_usage_limit(user_id)
    if not can_process:
        from config import FREE_USAGE_LIMIT
        limit_message = (
            "⚠️ Daily limit reached!\n\n"
            f"You've used all {FREE_USAGE_LIMIT} free document processing actions for today.\n\n"
            "💎 Upgrade to Premium for unlimited processing!\n"
            "Click /premium to see plans."
        )
        await message.reply(limit_message)
        return
    
    try:
        file_received_text = (
            "📄 Image received successfully!\n"
            "What would you like to do with it? 👇"
        )
        
        builder = InlineKeyboardBuilder()
        builder.button(text="📄 Convert to PDF", callback_data="image_to_pdf")
        builder.button(text="🗜️ Compress", callback_data="compress_image")
        builder.adjust(2)
        
        await message.reply(file_received_text, reply_markup=builder.as_markup())
        
        photo = message.photo[-1]
        await state.update_data(file_id=photo.file_id)
        await state.update_data(file_name="image.jpg")
        
    except Exception as e:
        logger.error(f"Error handling photo: {e}", exc_info=True)
        await message.reply(
            "⚠️ Oops! Something went wrong while processing your file.\n"
            "Please try again later or contact @DocuLunaSupport"
        )


async def handle_file_operation(callback: types.CallbackQuery, state: FSMContext):
    """Handle file operation callbacks with gamification integration."""
    operation = callback.data
    user_id = callback.from_user.id
    start_time = time.time()
    
    try:
        processing_text = (
            "⚙️ Please wait a moment...\n"
            "Your document is being processed 🔄"
        )
        
        await callback.message.edit_text(processing_text)
        await callback.answer()
        
        user_data = await state.get_data()
        file_id = user_data.get('file_id')
        file_name = user_data.get('file_name', 'file')
        
        result_file_path = None
        operation_type = None
        
        if operation == "convert_file":
            result_file_path = await convert_file(callback.bot, file_id, file_name, user_id)
            operation_type = "convert"
        elif operation == "compress_file":
            result_file_path = await compress_file(callback.bot, file_id, file_name, user_id)
            operation_type = "compress"
        elif operation == "compress_image":
            result_file_path = await compress_image(callback.bot, file_id, file_name)
            operation_type = "compress"
        elif operation == "image_to_pdf":
            result_file_path = await image_to_pdf(callback.bot, file_id, file_name, user_id)
            operation_type = "image_to_pdf"
        elif operation in ["merge_pdfs", "merge_pdf", "split_pdf", "split_pdfs"]:
            # Merge/split are handled in callbacks.py
            from handlers.callbacks import handle_merge_pdf_callback, handle_split_pdf_callback
            if operation in ["merge_pdfs", "merge_pdf"]:
                await handle_merge_pdf_callback(callback, state)
            else:
                await handle_split_pdf_callback(callback, state)
            return
        else:
            await callback.message.edit_text("Operation not yet implemented")
            return
        
        if result_file_path:
            # Calculate processing duration
            duration = time.time() - start_time
            file_size = os.path.getsize(result_file_path) if os.path.exists(result_file_path) else 0
            
            # Increment usage
            await increment_usage(user_id)
            
            # Process gamification rewards
            gamification_result = await process_gamification(
                user_id=user_id,
                operation=operation_type,
                filename=file_name,
                duration=duration,
                file_size=file_size,
                output_filename=os.path.basename(result_file_path)
            )
            
            from database.db import get_user_data
            from config import FREE_USAGE_LIMIT
            
            user_data = await get_user_data(user_id)
            usage_today = user_data.get('usage_today', 0) if user_data else 0
            is_premium = user_data.get('is_premium', False) if user_data else False
            
            # Build success message with gamification info
            from handlers.tool_instructions import format_file_size, get_operation_name
            
            success_text = "✅ Done! Your document is ready 🎉\n\n"
            
            # Add operation details
            success_text += f"📊 <b>{get_operation_name(operation_type)}</b>\n"
            success_text += f"├ File: {file_name}\n"
            duration = time.time() - start_time
            success_text += f"├ Time: {duration:.1f}s\n"
            success_text += f"└ Size: {format_file_size(file_size)}\n"
            
            # Add gamification rewards
            xp_gained = gamification_result.get("xp_gained", 0)
            if xp_gained > 0:
                success_text += f"\n⚡ +{xp_gained} XP earned!"
            
            if gamification_result.get("streak_increased"):
                streak = gamification_result.get('streak', 1)
                success_text += f" | 🔥 {streak} day streak!"
            
            # Show level up message
            if gamification_result.get("leveled_up"):
                new_level = gamification_result.get("new_level", 1)
                new_rank = gamification_result.get("new_rank", "")
                success_text += f"\n\n🎊 <b>Level Up!</b> You're now Level {new_level} - {new_rank}"
            
            if not is_premium:
                remaining = max(0, FREE_USAGE_LIMIT - usage_today)
                success_text += f"\n\n📊 <b>Usage Today:</b> {usage_today}/{FREE_USAGE_LIMIT} ({remaining} remaining)"
                
                if remaining == 0:
                    success_text += "\n💎 <i>Upgrade to Premium for unlimited!</i>"
                elif remaining == 1:
                    success_text += "\n⚠️ <i>Last free use for today!</i>"
            
            success_text += "\n\n<b>Download your file below 👇</b>"
            
            # Add gamification messages
            for msg in gamification_result.get("messages", []):
                success_text += f"\n\n{msg}"
            
            # Show achievement unlocks if any
            achievements = gamification_result.get("achievements", [])
            if achievements:
                from handlers.gamification import ACHIEVEMENT_MESSAGES
                for ach in achievements:
                    if ach in ACHIEVEMENT_MESSAGES:
                        success_text += f"\n\n🏆 {ACHIEVEMENT_MESSAGES[ach]}"
            
            # Get smart recommendation with exact format
            try:
                from handlers.smart_recommendation import smart_recommendation
                username = callback.from_user.first_name or "there"
                recommendation = await smart_recommendation.analyze_and_suggest(user_id)
                if recommendation and recommendation.get("message"):
                    action = recommendation['message'].split('\n')[0] if recommendation['message'] else "process more files"
                    smart_rec_msg = (
                        f"\n\n💡 {username}, your file's ready ✅\n"
                        f"I noticed this file could also be {action} to make your life easier ✨\n"
                        f"What's next? I'm still here 😎"
                    )
                    success_text += smart_rec_msg
            except Exception as e:
                logger.warning(f"Could not get recommendation: {e}")
            
            
            document = FSInputFile(result_file_path)
            await callback.message.answer_document(document)
            await callback.message.edit_text(success_text, parse_mode="HTML")
            
            # Show next action buttons
            try:
                builder = InlineKeyboardBuilder()
                if operation_type in ["convert", "pdf_to_word", "word_to_pdf"]:
                    builder.button(text="🗜️ Compress", callback_data="compress_pdf")
                    builder.button(text="↩️ Again", callback_data=operation_type)
                elif operation_type == "compress_pdf":
                    builder.button(text="📤 Share", url="https://t.me/share/url")
                    builder.button(text="🏠 Done", callback_data="back_to_menu")
                elif operation_type in ["merge_pdf"]:
                    builder.button(text="🗜️ Compress", callback_data="compress_pdf")
                    builder.button(text="🏠 Done", callback_data="back_to_menu")
                elif operation_type in ["split_pdf"]:
                    builder.button(text="↩️ Split Again", callback_data="split_pdf")
                    builder.button(text="🏠 Back", callback_data="back_to_menu")
                else:
                    builder.button(text="🔄 Another", callback_data="process_document")
                    builder.button(text="🏠 Menu", callback_data="back_to_menu")
                builder.adjust(2)
                await callback.message.answer(
                    "💡 <b>What's next?</b>\n\nChoose an action or process another file!",
                    reply_markup=builder.as_markup(),
                    parse_mode="HTML"
                )
            except:
                pass
            
            try:
                os.remove(result_file_path)
            except:
                pass
        else:
            # Log failed operation
            duration = time.time() - start_time
            await log_operation(
                user_id=user_id,
                operation=operation_type or operation,
                filename=file_name,
                duration=duration,
                status='failed'
            )
            raise Exception("Processing failed")
        
    except Exception as e:
        logger.error(f"Error in file operation {operation}: {e}", exc_info=True)
        
        error_msg = str(e).lower()
        
        if "size" in error_msg:
            from utils.messages import ERROR_OVERSIZED
            error_text = ERROR_OVERSIZED + "[  💎 Go Premium  ]  [  📤 Smaller File  ]  [  🏠 Back  ]"
        elif "format" in error_msg or "unsupported" in error_msg:
            from utils.messages import ERROR_UNSUPPORTED
            error_text = ERROR_UNSUPPORTED + "[  📤 Try Different  ]  [  ❓ Help  ]  [  🏠 Menu  ]"
        elif "corrupt" in error_msg or "invalid" in error_msg:
            from utils.messages import ERROR_CORRUPTED
            error_text = ERROR_CORRUPTED + "[  📤 Try Again  ]  [  🏠 Back  ]"
        elif "password" in error_msg or "protected" in error_msg:
            from utils.messages import ERROR_CORRUPTED_PDF
            error_text = ERROR_CORRUPTED_PDF + "[  📤 Send Unprotected  ]  [  🏠 Back  ]"
        elif "timeout" in error_msg or "slow" in error_msg:
            from utils.messages import ERROR_TIMEOUT
            error_text = ERROR_TIMEOUT + "[  🔄 Retry  ]  [  📤 Different  ]  [  🏠 Back  ]"
        else:
            error_text = "❌ Processing failed\n\nTry again or contact support if persistent.\n\n[  🔄 Retry  ]  [  🏠 Back  ]"
        
        await callback.message.edit_text(error_text)


async def convert_file(bot: Bot, file_id: str, file_name: str, user_id: int = None) -> str:
    """Convert PDF to Word or Word to PDF based on file type."""
    try:
        file = await bot.get_file(file_id)
        input_path = tempfile.mktemp(suffix=os.path.splitext(file_name)[1])
        await bot.download_file(file.file_path, input_path)
        
        if file_name.lower().endswith('.pdf'):
            output_path = tempfile.mktemp(suffix=".docx")
            result = PDFToWordConverter.convert_pdf_to_docx(
                input_path, 
                output_path, 
                preserve_layout=True, 
                extract_images=True
            )
            logger.info(f"PDF to Word conversion: {result['success']}")
            
            if user_id:
                from database.db import get_user_data
                user_data = await get_user_data(user_id)
                is_premium = user_data.get('is_premium', False) if user_data else False
                if not is_premium:
                    await add_watermark_to_docx(output_path)
                    logger.info(f"Added watermark to DOCX for free user {user_id}")
                    
        elif file_name.lower().endswith(('.docx', '.doc')):
            output_path = tempfile.mktemp(suffix=".pdf")
            result = WordToPDFConverter.convert_docx_to_pdf(
                input_path, 
                output_path, 
                preserve_formatting=True, 
                high_quality=True
            )
            logger.info(f"Word to PDF conversion: {result['success']}")
            
            if user_id:
                from database.db import get_user_data
                user_data = await get_user_data(user_id)
                is_premium = user_data.get('is_premium', False) if user_data else False
                if not is_premium:
                    await add_watermark_to_pdf(output_path)
                    logger.info(f"Added watermark to PDF for free user {user_id}")
                    
        else:
            raise ValueError("Unsupported file type for conversion")
        
        try:
            os.remove(input_path)
        except:
            pass
            
        return output_path
    except Exception as e:
        logger.error(f"Conversion error: {e}", exc_info=True)
        raise


async def compress_file(bot: Bot, file_id: str, file_name: str, user_id: int = None) -> str:
    """Compress PDF or DOCX file."""
    try:
        file = await bot.get_file(file_id)
        input_path = tempfile.mktemp(suffix=os.path.splitext(file_name)[1])
        await bot.download_file(file.file_path, input_path)
        
        output_path = tempfile.mktemp(suffix=os.path.splitext(file_name)[1])
        
        if file_name.lower().endswith('.pdf'):
            original_size, compressed_size = PDFCompressor.compress_pdf(
                input_path, 
                output_path, 
                CompressionLevel.MEDIUM
            )
            logger.info(f"PDF compressed: {original_size} → {compressed_size} bytes")
            
            if user_id:
                from database.db import get_user_data
                user_data = await get_user_data(user_id)
                is_premium = user_data.get('is_premium', False) if user_data else False
                if not is_premium:
                    await add_watermark_to_pdf(output_path)
                    logger.info(f"Added watermark to compressed PDF for free user {user_id}")
                    
        elif file_name.lower().endswith('.docx'):
            original_size, compressed_size = DOCXCompressor.compress_docx(
                input_path, 
                output_path, 
                CompressionLevel.MEDIUM
            )
            logger.info(f"DOCX compressed: {original_size} → {compressed_size} bytes")
            
            if user_id:
                from database.db import get_user_data
                user_data = await get_user_data(user_id)
                is_premium = user_data.get('is_premium', False) if user_data else False
                if not is_premium:
                    await add_watermark_to_docx(output_path)
                    logger.info(f"Added watermark to compressed DOCX for free user {user_id}")
                    
        else:
            raise ValueError("Unsupported file type for compression")
        
        try:
            os.remove(input_path)
        except:
            pass
            
        return output_path
    except Exception as e:
        logger.error(f"Compression error: {e}", exc_info=True)
        raise


async def compress_image(bot: Bot, file_id: str, file_name: str) -> str:
    """Compress image file."""
    try:
        file = await bot.get_file(file_id)
        input_path = tempfile.mktemp(suffix=".jpg")
        await bot.download_file(file.file_path, input_path)
        
        output_path = tempfile.mktemp(suffix=".jpg")
        
        # Open and compress image
        with Image.open(input_path) as img:
            # Convert to RGB if needed
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Save with compression
            img.save(output_path, 'JPEG', quality=85, optimize=True)
        
        original_size = os.path.getsize(input_path)
        compressed_size = os.path.getsize(output_path)
        logger.info(f"Image compressed: {original_size} → {compressed_size} bytes")
        
        try:
            os.remove(input_path)
        except:
            pass
            
        return output_path
    except Exception as e:
        logger.error(f"Image compression error: {e}", exc_info=True)
        raise


async def image_to_pdf(bot: Bot, file_id: str, file_name: str, user_id: int = None) -> str:
    """Convert image to PDF."""
    try:
        file = await bot.get_file(file_id)
        input_path = tempfile.mktemp(suffix=".jpg")
        await bot.download_file(file.file_path, input_path)
        
        output_path = tempfile.mktemp(suffix=".pdf")
        
        image = Image.open(input_path)
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        img_width, img_height = image.size
        c = canvas.Canvas(output_path, pagesize=A4)
        page_width, page_height = A4
        
        scale = min(page_width / img_width, page_height / img_height) * 0.9
        x = (page_width - img_width * scale) / 2
        y = (page_height - img_height * scale) / 2
        
        c.drawImage(input_path, x, y, img_width * scale, img_height * scale)
        c.save()
        
        if user_id:
            from database.db import get_user_data
            user_data = await get_user_data(user_id)
            is_premium = user_data.get('is_premium', False) if user_data else False
            if not is_premium:
                await add_watermark_to_pdf(output_path)
                logger.info(f"Added watermark to PDF for free user {user_id}")
        
        try:
            os.remove(input_path)
        except:
            pass
            
        logger.info(f"Image to PDF conversion complete: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Image to PDF error: {e}", exc_info=True)
        raise


async def add_watermark_to_pdf(pdf_path: str):
    """Add watermark to PDF using comprehensive watermark system."""
    try:
        # Read PDF file
        with open(pdf_path, 'rb') as f:
            pdf_data = f.read()
        
        # Configure watermark for bottom of page
        config = WatermarkConfig(
            text=WATERMARK_TEXT,
            position="bottom-center",
            opacity=0.4,
            font_size=9,
            font_color=(128, 128, 128),
            rotation=0
        )
        
        # Add watermark
        watermarked_data = await watermark_manager.add_text_watermark(
            BytesIO(pdf_data),
            WATERMARK_TEXT,
            config,
            output_format='.pdf'
        )
        
        # Save watermarked PDF
        with open(pdf_path, 'wb') as f:
            f.write(watermarked_data)
        
        logger.info(f"Added watermark to PDF: {pdf_path}")
        
    except Exception as e:
        logger.error(f"Error adding PDF watermark: {e}", exc_info=True)


async def add_watermark_to_docx(docx_path: str):
    """Add watermark to DOCX using footer."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Run in executor since python-docx is sync
        loop = asyncio.get_event_loop()
        
        def add_footer_watermark():
            doc = Document(docx_path)
            
            # Add to footer of first section
            section = doc.sections[0]
            footer = section.footer
            
            if footer.paragraphs:
                paragraph = footer.paragraphs[0]
            else:
                paragraph = footer.add_paragraph()
            
            paragraph.text = WATERMARK_TEXT
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if paragraph.runs:
                run = paragraph.runs[0]
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.italic = True
            
            doc.save(docx_path)
        
        await loop.run_in_executor(None, add_footer_watermark)
        logger.info(f"Added watermark to DOCX: {docx_path}")
        
    except Exception as e:
        logger.error(f"Error adding DOCX watermark: {e}", exc_info=True)


def register_file_handlers(dp: Dispatcher):
    """Register file handlers."""
    dp.message.register(handle_document, lambda m: m.document is not None)
    dp.message.register(handle_photo, lambda m: m.photo is not None)
    dp.callback_query.register(handle_file_operation, lambda c: c.data in [
        "convert_file", "merge_pdfs", "split_pdf", "compress_file", 
        "compress_image", "image_to_pdf"
    ])
    logger.info("✓ File handlers registered with gamification integration")
